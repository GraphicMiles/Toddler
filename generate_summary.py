from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_summary():
    doc = Document()

    # Title
    title = doc.add_heading('Toddler Project Summary & Roadmap', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project Overview
    doc.add_heading('Project Overview', level=1)
    doc.add_paragraph(
        "Toddler is a professional-grade, no-code AI training platform designed for domain experts. "
        "It enables non-technical users to upload their own datasets (CSVs), train specialized "
        "classifiers in seconds, and own their model artifacts forever."
    )

    # Technologies Used
    doc.add_heading('Technologies Used', level=1)
    techs = [
        ("Frontend", "React (v19), Vite, Tailwind CSS v4, Lucide Icons, React Router."),
        ("Backend", "FastAPI (Python 3.9), Scikit-learn (ML Engine), Pandas (Data Processing)."),
        ("Infrastructure", "Firebase Auth (Authentication), Firestore (Metadata/NoSQL), Render (API Hosting), Vercel (Frontend Hosting)."),
        ("Design", "Space Grotesk & Inter typography, Neon Dark V2 Theme (Violet-Lime).")
    ]
    for category, desc in techs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{category}: ")
        run.bold = True
        p.add_run(desc)

    # Features Implemented
    doc.add_heading('Features Implemented', level=1)
    features = [
        ("Award-Tier Landing Page", "High-converting, responsive design with native CSS animations, premium film-grain overlays, and ambient gradient orbs."),
        ("Direct-to-Backend Training", "Secure CSV ingestion that bypasses standard storage limits for the free tier, processing data directly in synchronous cycles."),
        ("Professional Auth Suite", "Split-screen Login and Signup with Google/Apple SSO support and secure form validation."),
        ("Multi-step Onboarding", "Guided flow for project naming, dataset ingestion, and column mapping (Text vs. Label)."),
        ("Intelligence Dashboard", "Centralized hub with left-aligned technical manual aesthetic and Skeleton Shimmer loading states."),
        ("Interactive Playground", "Logic terminal with real-time word-level explainability highlights based on model weights."),
        ("Analytics Suite", "Interactive Confusion Matrices, Global Feature Importance charts, and Dataset Balance metrics."),
        ("Intent-Based Chatbot", "A manager to map training labels to custom responses and a live chat preview window."),
        ("Model Portability", "One-click export of .pkl (Scikit-learn) artifacts, fulfilling the promise of 100% data ownership."),
        ("Developer Hub", "API Key generation and Python integration snippets for local model execution.")
    ]
    for feat, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{feat}: ")
        run.bold = True
        p.add_run(desc)

    # Future Roadmap (Unicorn Pass)
    doc.add_heading('Future Roadmap (The 30 Upgrades)', level=1)
    roadmap = [
        "AutoML Tournament: Automatically testing multiple ML algorithms (XGBoost, Random Forest).",
        "Transformer Embeddings: Deep semantic understanding using small language models.",
        "Active Learning: Surfacing low-confidence rows for manual label correction.",
        "Embeddable Chat Widget: JS snippet for hosting Toddler bots on external sites.",
        "Webhook Support: Real-time triggers for automated actions via Zapier/Make.",
        "Team Workspaces: Multi-user collaboration on datasets and models.",
        "SOC2 Compliance Export: Standardized security reporting for enterprise users.",
        "Chrome Extension: Classifying text directly from any web browser window.",
        "Python SDK: Native library for developers to import Toddler models into local code.",
        "Model Versioning Timeline: Historical records and rollbacks for trained artifacts."
    ]
    for item in roadmap:
        doc.add_paragraph(item, style='List Number')

    doc.save('toddler_project_summary.docx')

if __name__ == "__main__":
    create_summary()
